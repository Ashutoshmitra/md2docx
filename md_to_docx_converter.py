"""
Markdown to Word Converter
Core conversion functionality for transforming Markdown to Word with template support.
"""

import os
import re
from pathlib import Path
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches

class MarkdownToWordConverter:
    """
    Converts Markdown files to Word documents using a specified template.
    Handles headings, lists, and special formatting requirements.
    """
    
    def __init__(self, template_path=None):
        """
        Initialize the converter with an optional template.
        
        Args:
            template_path (str): Path to the Word template (.dotx) file
        """
        self.template_path = template_path
        
    def convert_file(self, input_path, output_path=None):
        """
        Convert a single Markdown file to Word format.
        
        Args:
            input_path (str): Path to the input Markdown file
            output_path (str, optional): Path for the output Word file
            
        Returns:
            str: Path to the generated Word document
        """
        # Determine output path if not specified
        if not output_path:
            output_dir = os.path.dirname(input_path)
            filename = os.path.basename(input_path)
            base_name = os.path.splitext(filename)[0]
            output_path = os.path.join(output_dir, f"{base_name}.docx")
            
        # Read markdown content
        with open(input_path, 'r', encoding='utf-8') as md_file:
            md_content = md_file.read()
            
        # Convert to Word
        self._convert_content(md_content, output_path)
        
        return output_path
    
    def convert_directory(self, input_dir, output_dir=None):
        """
        Convert all Markdown files in a directory to Word format.
        
        Args:
            input_dir (str): Path to the input directory
            output_dir (str, optional): Path for the output directory
            
        Returns:
            list: Paths to all generated Word documents
        """
        # Use input directory as output if not specified
        if not output_dir:
            output_dir = input_dir
            
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        output_files = []
        
        # Process each markdown file
        for filename in os.listdir(input_dir):
            if filename.lower().endswith('.md'):
                input_path = os.path.join(input_dir, filename)
                base_name = os.path.splitext(filename)[0]
                output_path = os.path.join(output_dir, f"{base_name}.docx")
                
                output_files.append(self.convert_file(input_path, output_path))
                
        return output_files
    
    def _convert_content(self, md_content, output_path):
        """
        Convert markdown content to a Word document.
        
        Args:
            md_content (str): Markdown content as string
            output_path (str): Path for the output Word file
        """
        # Create document from template if specified, otherwise create new
        if self.template_path and os.path.exists(self.template_path):
            doc = Document(self.template_path)
        else:
            doc = Document()
            
        # Convert markdown to HTML
        html_content = markdown.markdown(
            md_content,
            extensions=['markdown.extensions.tables', 'markdown.extensions.fenced_code']
        )
        
        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Process content and apply styles
        self._process_html_to_docx(soup, doc)
        
        # Save the document
        doc.save(output_path)
    
    def _process_html_to_docx(self, soup, doc):
        """
        Process HTML content and convert it to Word with appropriate styling.
        
        Args:
            soup (BeautifulSoup): Parsed HTML content
            doc (Document): Word document object
        """
        # Remove existing paragraphs from template (if any)
        if len(doc.paragraphs) > 0 and not doc.paragraphs[0].text.strip():
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)
            p._p = p._element = None
        
        for element in soup.children:
            if element.name is None:
                continue
                
            # Process headings
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                self._process_heading(element, doc)
                
            # Process paragraphs
            elif element.name == 'p':
                self._process_paragraph(element, doc)
                
            # Process lists
            elif element.name in ['ul', 'ol']:
                self._process_list(element, doc)
                
            # Process code blocks
            elif element.name == 'pre':
                self._process_code_block(element, doc)
                
            # Process tables
            elif element.name == 'table':
                self._process_table(element, doc)
    
    def _process_heading(self, element, doc):
        """Process heading elements."""
        level = int(element.name[1])  # Extract heading level (1-6)
        
        # Determine the corresponding Word heading style
        style_name = f'Heading {level}'
        
        # Create paragraph with appropriate style
        p = doc.add_paragraph()
        p.style = style_name
        
        # Remove existing numbering (if any)
        text = element.get_text()
        text = re.sub(r'^\d+(\.\d+)*\s+', '', text)
        
        # Add the text
        p.add_run(text)
    
    def _process_paragraph(self, element, doc):
        """Process paragraph elements."""
        # Skip empty paragraphs
        if not element.get_text().strip():
            return
            
        # Create paragraph
        p = doc.add_paragraph()
        p.style = 'Normal'
        
        # Process all child elements (bold, italic, etc.)
        self._process_inline_elements(element, p)
        
        # Remove soft returns (replace with spaces)
        for run in p.runs:
            run.text = run.text.replace('\n', ' ')
    
    def _process_inline_elements(self, element, paragraph):
        """Process inline formatting elements within a paragraph."""
        for child in element.children:
            if child.name is None:  # Regular text
                run = paragraph.add_run(child.string)
            elif child.name == 'strong' or child.name == 'b':
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name == 'em' or child.name == 'i':
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == 'code':
                run = paragraph.add_run(child.get_text())
                run.font.name = 'Courier New'
            elif child.name == 'a':
                run = paragraph.add_run(child.get_text())
                if child.get('href'):
                    run.hyperlink = child.get('href')
            else:
                # Handle other elements recursively
                self._process_inline_elements(child, paragraph)
    
    def _process_list(self, element, doc, level=0):
        """Process list elements (ordered and unordered)."""
        is_ordered = element.name == 'ol'
        
        for item in element.find_all('li', recursive=False):
            # Create paragraph for list item
            p = doc.add_paragraph()
            
            # Determine list style
            if is_ordered:
                style_name = 'List Number'
            else:
                style_name = 'List Bullet'
                
            p.style = style_name
            p.paragraph_format.left_indent = Inches(level * 0.25)
            
            # Process all content of list item (except nested lists)
            # First, let's get direct text content
            for child in item.children:
                if child.name not in ['ul', 'ol']:
                    if child.name is None:  # Direct text node
                        if child.string and child.string.strip():
                            p.add_run(child.string.strip())
                    else:  # Other elements like <strong>, <em>, etc.
                        self._process_inline_elements(child, p)
            
            # Process nested lists recursively
            nested_lists = item.find_all(['ul', 'ol'], recursive=False)
            for nested_list in nested_lists:
                self._process_list(nested_list, doc, level + 1)
    
    def _process_code_block(self, element, doc):
        """Process code blocks with proper line breaks and formatting."""
        # Extract code content
        code = element.get_text()
        
        # Create paragraph for code block
        p = doc.add_paragraph()
        p.style = 'No Spacing'
        
        # Split code by lines and add each line separately to preserve formatting
        lines = code.split('\n')
        for i, line in enumerate(lines):
            if i > 0:  # Add line break before each line (except the first one)
                p.add_run('\n')
            # Add the line with preserved spaces
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
    
    def _process_table(self, element, doc):
        """Process table elements."""
        # Count rows and columns
        rows = element.find_all('tr')
        if not rows:
            return
            
        max_cols = 0
        for row in rows:
            cols = row.find_all(['td', 'th'])
            max_cols = max(max_cols, len(cols))
            
        if max_cols == 0:
            return
            
        # Create table
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        # Fill table
        for r_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for c_idx, cell in enumerate(cells):
                if c_idx < max_cols:  # Stay within bounds
                    # Set cell content
                    text = cell.get_text().strip()
                    table.cell(r_idx, c_idx).text = text
                    
                    # Apply header formatting
                    if cell.name == 'th':
                        for paragraph in table.cell(r_idx, c_idx).paragraphs:
                            for run in paragraph.runs:
                                run.bold = True