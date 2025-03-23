import os
import re
from docx import Document
import markdown
from bs4 import BeautifulSoup
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx_from_readme(readme_path, template_path, output_path):
    # Load the template document
    doc = Document(template_path)

    # Dictionary to track list numbering for different sections
    list_counters = {}
    # Create a unique identifier for the current list section
    current_list_section = 0

    # Preserve headers and footers by removing only body paragraphs
    body = doc._body
    for element in body._element[:]:
        if element.tag.endswith('p'):
            body._element.remove(element)

    # Read the README file and remove soft carriage returns
    with open(readme_path, 'r', encoding='utf-8') as f:
        readme_content = f.read().replace('\r\n', '\n').replace('\r', '\n')

    # Convert Markdown to HTML
    html_content = markdown.markdown(readme_content, extensions=['extra', 'tables', 'fenced_code'])
    soup = BeautifulSoup(html_content, 'html.parser')

    # Helper function to clean manual numbering from headings
    def clean_heading_text(text):
        # First pattern for main headings like "1. TITLE"
        main_pattern = r'^(\d+)\.\s+'
        # Second pattern for sub-headings like "1.1 Title"
        sub_pattern = r'^(\d+\.\d+)\s+'
        
        # Try main heading pattern first
        if re.match(main_pattern, text):
            cleaned = re.sub(main_pattern, '', text)
        # Then try sub-heading pattern
        elif re.match(sub_pattern, text):
            cleaned = re.sub(sub_pattern, '', text)
        else:
            cleaned = text
            
        return cleaned

    # Helper function to add styled paragraph with bold support
    def add_paragraph(doc, element, style_name):
        paragraph = doc.add_paragraph()
        try:
            paragraph.style = doc.styles[style_name]
        except KeyError:
            print(f"Warning: Style '{style_name}' not found in template. Using 'Normal'.")
            paragraph.style = doc.styles['Normal']
        
        # Get and clean text for headings
        raw_text = element.get_text().strip()
        if style_name.startswith('Heading'):
            text = clean_heading_text(raw_text)
        else:
            text = raw_text
        
        # Process contents with bold support
        if element.find('strong'):
            for content in element.contents:
                if isinstance(content, str):
                    if content.strip():  # Only add non-empty strings
                        paragraph.add_run(content.strip())
                elif content.name == 'strong':
                    run = paragraph.add_run(content.get_text().strip())
                    run.bold = True
        else:
            paragraph.add_run(text)
        
        return paragraph

    # Helper function to add code block
    def add_code_block(doc, code_text, language=None):
        container = doc.add_paragraph()
        container.style = doc.styles['Normal']
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), "F0F0F0")
        container._p.get_or_add_pPr().append(shading_elm)
        for line in code_text.strip().split('\n'):
            if line.strip() == '':
                container.add_run('\n')
            else:
                run = container.add_run(line + '\n')
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        return container

    # Helper function to add table
    def add_table(doc, table_soup):
        rows = table_soup.find_all('tr')
        if not rows:
            return
        max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    cell_text = cell.get_text().strip()
                    paragraph = table.rows[i].cells[j].paragraphs[0]
                    if cell.name == 'th':
                        run = paragraph.add_run(cell_text)
                        run.bold = True
                    else:
                        paragraph.add_run(cell_text)
    
    # Add a dash-style bullet list item
    def add_bullet_item(doc, text):
        para = doc.add_paragraph(style="–bullet 1")
        para.add_run(text)
        return para
        
    # Add a dash-style bullet for a nested level
    def add_nested_bullet_item(doc, text, level):
        style_name = f"–bullet {level}" if level <= 4 else "–bullet 4"
        try:
            para = doc.add_paragraph(style=style_name)
        except KeyError:
            para = doc.add_paragraph(style="–bullet 1")
        para.add_run(text)
        return para
    
    # Process lists hierarchically
    def process_lists(doc, element, is_root=True):
        nonlocal current_list_section
        
        # Only increment section counter at root level
        if is_root:
            current_list_section += 1
        
        if element.name == 'ol':
            # Initialize counter for this list section if it doesn't exist
            if current_list_section not in list_counters:
                list_counters[current_list_section] = 0
                
            # Process ordered list items
            items = element.find_all('li', recursive=False)
            for item in items:
                # Extract text content
                text_content = ''
                for content in item.contents:
                    if isinstance(content, str):
                        text_content += content
                    elif content.name not in ['ul', 'ol']:
                        text_content += content.get_text()
                
                # Create a paragraph with appropriate numbering
                para = doc.add_paragraph(style="| Text numbering 1")
                
                # This is a hack to reset numbering between list sections
                para._p.get_or_add_pPr().append(OxmlElement('w:numPr'))
                para.add_run(text_content.strip())
                
                # Process nested lists
                for child in item.children:
                    if hasattr(child, 'name'):
                        if child.name in ['ul', 'ol']:
                            process_lists(doc, child, False)
        
        elif element.name == 'ul':
            # Process unordered list items
            items = element.find_all('li', recursive=False)
            for item in items:
                # Extract text content
                text_content = ''
                for content in item.contents:
                    if isinstance(content, str):
                        text_content += content
                    elif content.name not in ['ul', 'ol']:
                        text_content += content.get_text()
                
                # Create a bullet point
                para = add_bullet_item(doc, text_content.strip())
                
                # Process nested lists
                for child in item.children:
                    if hasattr(child, 'name'):
                        if child.name in ['ul', 'ol']:
                            process_lists(doc, child, False)

    # Function to create content like in the original document format
    # For dash-style bullets using "-" 
    def format_dash_list(text):
        if not text.startswith('-'):
            return text
        
        items = text.split('\n- ')
        result = items[0] + '\n'
        for item in items[1:]:
            result += f"- {item}\n"
        return result

    # Process HTML elements
    elements = [element for element in soup if hasattr(element, 'name') and element.name]
    
    for element in elements:
        if element.name == 'h1':
            add_paragraph(doc, element, 'Heading 1')
        elif element.name == 'h2':
            add_paragraph(doc, element, 'Heading 2')
        elif element.name == 'h3':
            add_paragraph(doc, element, 'Heading 3')
        elif element.name == 'h4':
            add_paragraph(doc, element, 'Heading 4')
        elif element.name == 'h5':
            add_paragraph(doc, element, 'Heading 5')
        elif element.name == 'p':
            if element.find('code'):
                code_text = element.find('code').get_text()
                add_code_block(doc, code_text)
            else:
                # Check if this paragraph contains a "- " list syntax for original document format
                text = element.get_text()
                if text.strip().startswith('-'):
                    # This is a special case for handling dash lists
                    lines = text.strip().split('\n')
                    for line in lines:
                        if line.strip().startswith('-'):
                            p = doc.add_paragraph(style="–bullet 1")
                            p.add_run(line.strip()[2:].strip())  # Remove the dash and add content
                        else:
                            add_paragraph(doc, BeautifulSoup(f"<p>{line}</p>", 'html.parser').p, 'Normal')
                else:
                    add_paragraph(doc, element, 'Normal')
        elif element.name == 'pre':
            code_block = element.find('code')
            if code_block:
                language = None
                if 'class' in code_block.attrs:
                    language_classes = [cls for cls in code_block['class'] if cls.startswith('language-')]
                    if language_classes:
                        language = language_classes[0].replace('language-', '')
                code_text = code_block.get_text().replace('\r\n', '\n').replace('\r', '\n')
                add_code_block(doc, code_text, language)
        elif element.name == 'ul':
            process_lists(doc, element)
        elif element.name == 'ol':
            process_lists(doc, element)
        elif element.name == 'table':
            add_table(doc, element)

    # Save the document
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    readme_file = "/Users/ashutoshmitra/md2word/briefing-document.md"
    template_file = "/Users/ashutoshmitra/md2word/samples/Template - Portrait.docx"
    output_file = "/Users/ashutoshmitra/md2word/briefing-document.docx"

    if not os.path.exists(readme_file):
        print(f"Error: {readme_file} not found.")
    elif not os.path.exists(template_file):
        print(f"Error: {template_file} not found.")
    else:
        create_docx_from_readme(readme_file, template_file, output_file)