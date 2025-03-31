#!/usr/bin/env python3
"""
MD to DOCX Converter Module (Enhanced Version)
Converts Markdown files to Word document format using a specified template
with style extraction for consistent formatting
"""

import os
import re
from docx import Document
import markdown
from bs4 import BeautifulSoup
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_template_styles(doc):
    """
    Extract style information from a Word template document.
    Modify to skip certain elements like initial tables or headers.
    
    Args:
        doc: A python-docx Document object
        
    Returns:
        dict: Dictionary containing style information
    """
    # Structure to hold style information
    styles_info = {
        "paragraph_styles": {},
        "character_styles": {},
        "table_styles": {},
        "numbering_styles": {},
        "bullet_styles": {},
        "default_bullets": ['•', '◦', '▪', '▫']
    }
    
    # Only process styles from body
    body_elements = [el for el in doc._body._element if not el.tag.endswith('tbl')]
    
    # Extract styles from body paragraphs
    for style in doc.styles:
        try:
            if style.type == 1:  # Paragraph style
                para_style_info = {
                    "name": style.name,
                    "font_name": style.font.name if hasattr(style.font, 'name') else None,
                    "font_size": style.font.size.pt if hasattr(style.font, 'size') and style.font.size else None,
                    "bold": style.font.bold if hasattr(style.font, 'bold') else None,
                    "italic": style.font.italic if hasattr(style.font, 'italic') else None,
                    "underline": style.font.underline if hasattr(style.font, 'underline') else None
                }
                styles_info['paragraph_styles'][style.name] = para_style_info
                
                # Identify list styles
                if any(marker in style.name.lower() for marker in ['bullet', '•', '◦', '▪', '▫']):
                    styles_info['bullet_styles'][style.name] = para_style_info
                
                # Identify numbering styles
                if any(marker in style.name.lower() for marker in ['number', '1.', 'a.', 'text numbering']):
                    styles_info['numbering_styles'][style.name] = para_style_info
        except Exception as e:
            print(f"Error processing style {style.name}: {e}")
        # Add table style extraction

    table_styles = extract_table_styles(doc)
    styles_info['table_style_details'] = table_styles

    return styles_info

def find_list_formats(doc, styles_info):
    """
    Analyze paragraphs in the document to find bullet and numbering formats
    
    Args:
        doc: A python-docx Document object
        styles_info: Dictionary to update with list format information
    """
    # Map to track bullet characters and numbering formats by style
    bullet_formats = {}
    numbering_formats = {}
    
    # Process all paragraphs to find list formats
    for para in doc.paragraphs:
        # Skip if not a paragraph with content
        if not para.text.strip():
            continue
            
        style_name = para.style.name if hasattr(para, 'style') and hasattr(para.style, 'name') else 'Normal'
        
        # Check if this is a bullet list paragraph
        if para.text.strip().startswith('•') or para.text.strip().startswith('◦') or para.text.strip().startswith('▪'):
            bullet_char = para.text.strip()[0]
            if style_name not in bullet_formats:
                bullet_formats[style_name] = bullet_char
        
        # Check if this is a numbered list paragraph (e.g., "1. Text")
        elif para.text.strip() and para.text.strip()[0].isdigit() and ". " in para.text.strip()[:4]:
            if style_name not in numbering_formats:
                numbering_formats[style_name] = "decimal"
        
        # Check if this is a lettered list paragraph (e.g., "a. Text")
        elif para.text.strip() and para.text.strip()[0].isalpha() and ". " in para.text.strip()[:4]:
            if style_name not in numbering_formats:
                numbering_formats[style_name] = "lowerLetter"
    
    # Update styles_info with the discovered formats
    styles_info['bullet_formats'] = bullet_formats
    styles_info['numbering_formats'] = numbering_formats
    
    # Fallback bullet characters if none found in the document
    if not bullet_formats:
        styles_info['default_bullets'] = ['•', '◦', '▪', '▫']
    else:
        styles_info['default_bullets'] = list(bullet_formats.values())
        
    # Ensure we have at least the default bullet characters
    if not styles_info['default_bullets']:
        styles_info['default_bullets'] = ['•', '◦', '▪', '▫']

def extract_table_styles(doc):
    """
    Enhanced table style extraction from template document.
    
    Args:
        doc: python-docx Document object
    
    Returns:
        dict: Detailed table style information
    """
    table_styles = {}
    
    # Process all tables in the template
    for table_index, table in enumerate(doc.tables):
        style_name = table.style.name if hasattr(table, 'style') and hasattr(table.style, 'name') else 'Default Table Style'
        
        # Initialize style dictionary for this table style
        table_styles[style_name] = {
            "name": style_name,
            "has_header": False,
            "header_style": "Table Heading",
            "header_formatting": {},
            "cell_styles": {},
            "table_index": table_index,  # Store the table index for reference
        }
        
        # Check for header row
        if len(table.rows) > 1:
            first_row = table.rows[0]
            second_row = table.rows[1]
            
            # Detect header by formatting differences
            header_cells = first_row.cells
            body_cells = second_row.cells
            
            # Check for bold text in header
            header_bold = any(
                para.runs and para.runs[0].bold 
                for cell in header_cells 
                for para in cell.paragraphs
            )
            
            # Check for color or other formatting differences
            if header_bold:
                table_styles[style_name]["has_header"] = True
                table_styles[style_name]["header_formatting"] = {
                    "bold": True
                }
                
            # Capture cell text style (for first cell in second row)
            if len(body_cells) > 0 and len(body_cells[0].paragraphs) > 0:
                first_body_cell = body_cells[0]
                cell_para = first_body_cell.paragraphs[0]
                
                if hasattr(cell_para, 'style') and hasattr(cell_para.style, 'name'):
                    table_styles[style_name]["cell_style"] = cell_para.style.name
                    
                # Capture run properties if any
                if len(cell_para.runs) > 0:
                    run = cell_para.runs[0]
                    table_styles[style_name]["cell_font"] = {
                        "name": run.font.name if hasattr(run.font, 'name') else None,
                        "size": run.font.size.pt if hasattr(run.font, 'size') and run.font.size else None,
                        "bold": run.font.bold if hasattr(run.font, 'bold') else None,
                        "italic": run.font.italic if hasattr(run.font, 'italic') else None,
                    }
    
    # If we found any tables, store additional information about the first table (usually logo table)
    if doc.tables and len(doc.tables) > 0:
        first_table = doc.tables[0]
        table_styles["first_table"] = {
            "row_count": len(first_table.rows),
            "col_count": len(first_table.rows[0].cells) if len(first_table.rows) > 0 else 0,
        }
    
    return table_styles

def clean_heading_text(text):
    """
    Clean manual numbering from headings.
    
    Args:
        text (str): The heading text
        
    Returns:
        str: Cleaned heading text
    """
    # Match patterns like "1.", "1.1", "1.1.1" at the start of a heading
    # Updated to be more aggressive in removing all section numbers
    pattern = r'^(\d+(\.\d+)*(\s+|\.|\s+\.)?)\s*'
    if re.match(pattern, text):
        return re.sub(pattern, '', text)
    return text

def add_paragraph(doc, element, style_name):
    """
    Add a styled paragraph with formatting support.
    
    Args:
        doc: The docx Document object
        element: BeautifulSoup element containing the paragraph content
        style_name: The style name to apply
        
    Returns:
        The created paragraph object
    """
    paragraph = doc.add_paragraph()
    
    # Try to set the style
    try:
        paragraph.style = style_name
    except KeyError:
        print(f"Warning: Style '{style_name}' not found in template. Using 'Normal'.")
        paragraph.style = 'Normal'
    
    # Get and clean text for headings
    raw_text = element.get_text().strip()
    if style_name.startswith('Heading'):
        text = clean_heading_text(raw_text)
    else:
        text = raw_text
    
    # Process contents with formatting support
    if element.find(['strong', 'em', 'i', 'b', 'code']):
        # Handle paragraph with formatted text
        for content in element.children:
            if isinstance(content, str):
                paragraph.add_run(content)
            elif content.name == 'strong' or content.name == 'b':
                run = paragraph.add_run(content.get_text())
                run.bold = True
            elif content.name == 'em' or content.name == 'i':
                run = paragraph.add_run(content.get_text())
                run.italic = True
            elif content.name == 'code':
                run = paragraph.add_run(content.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
    else:
        # Simple paragraph without formatting
        paragraph.add_run(text)
    
    return paragraph

def add_code_block(doc, code_text, language=None):
    """
    Add a code block to the document.
    
    Args:
        doc: The docx Document object
        code_text: The code text to add
        language: Optional language specification
        
    Returns:
        The created paragraph object
    """
    container = doc.add_paragraph()
    container.style = 'Normal'
    
    # Add gray background
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), "F0F0F0")
    container._p.get_or_add_pPr().append(shading_elm)
    
    # Add code content line by line with courier font
    lines = code_text.strip().split('\n')
    for i, line in enumerate(lines):
        if line.strip() == '':
            container.add_run('\n')
        else:
            run = container.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            
            # Add newline if not the last line
            if i < len(lines) - 1:
                container.add_run('\n')
    
    return container

def add_table_with_styles(doc, table_soup, styles_info):
    """
    Add a table to the document using template styles with proper style inheritance.
    
    Args:
        doc: The docx Document object
        table_soup: BeautifulSoup object for the table
        styles_info: Dictionary containing style information from the template
    """
    # Extract rows from the table
    rows = table_soup.find_all('tr')
    if not rows:
        return
    
    # Check if this is a meaningful data table (more than just headers)
    if len(rows) <= 1:
        return
    
    # Try to ensure this is a data table with actual content
    content_rows = [row for row in rows[1:] if row.get_text(strip=True)]
    if not content_rows:
        return
    
    # Determine maximum number of columns
    max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
    
    # Create the table in the document
    table = doc.add_table(rows=len(rows), cols=max_cols)
    
    # Get detailed table style information from template
    table_style_details = styles_info.get('table_style_details', {})
    
    # Find the most appropriate table style from the template - prefer the second table if available
    # (assuming first table is usually the logo table)
    preferred_table_style = None
    
    # If there are at least two tables in the template, use the second one's style
    for style_name, details in table_style_details.items():
        if details.get("table_index", 0) == 1:  # Second table (index 1)
            preferred_table_style = style_name
            break
    
    # If no second table, fall back to any other table style
    if not preferred_table_style:
        for style_name, details in table_style_details.items():
            if style_name != "first_table":
                preferred_table_style = style_name
                break
    
    # Apply table style - use Table Grid as a fallback
    default_table_style = preferred_table_style or 'Table Grid'
    try:
        table.style = default_table_style
        print(f"Applied table style: {default_table_style}")
    except Exception as e:
        print(f"Could not apply table style {default_table_style}, falling back to Table Grid: {e}")
        try:
            table.style = 'Table Grid'
        except:
            pass
    
    # Get cell styles from template
    cell_style_name = "= Table text"  # Default style name for regular cells
    header_style_name = "Table Heading"  # Default style name for header cells
    
    # Check if these styles exist in the document
    if cell_style_name not in [s.name for s in doc.styles]:
        available_table_styles = [s.name for s in doc.styles if '= Table' in s.name or 'Table' in s.name]
        if available_table_styles:
            cell_style_name = available_table_styles[0]
        else:
            cell_style_name = "Normal"
            
    # Check if Table Heading style exists
    if header_style_name not in [s.name for s in doc.styles]:
        available_header_styles = [s.name for s in doc.styles if 'Table Head' in s.name]
        if available_header_styles:
            header_style_name = available_header_styles[0]
        else:
            header_style_name = "Normal"
    
    # Determine if table has a header
    has_header = any(row.find_all('th') for row in rows) or True  # Assume first row is header if no th found
    
    # Process each row and cell
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        
        for j, cell in enumerate(cells):
            if j >= max_cols:
                break
            
            # Get cell text
            cell_text = cell.get_text().strip()
            table_cell = table.rows[i].cells[j]
            
            # Clear any existing paragraphs in the cell
            for p in table_cell.paragraphs[:]:
                p._element.getparent().remove(p._element)
            
            # Add a new paragraph to the cell
            paragraph = table_cell.add_paragraph()
            
            # Apply appropriate style based on whether this is a header cell
            if (cell.name == 'th') or (i == 0 and has_header):
                try:
                    paragraph.style = header_style_name
                except:
                    # If style application fails, just continue with default
                    pass
                
                # Add content - keeping bold for headers regardless of style
                run = paragraph.add_run(cell_text)
                run.bold = True
            else:
                # Apply the table text style for regular cells
                try:
                    paragraph.style = cell_style_name
                except:
                    # If style application fails, just continue with default
                    pass
                
                # Simply add the text without any direct formatting
                run = paragraph.add_run(cell_text)
    
    # If the table has a special format for the bottom row (totals)
    if len(rows) > 1:
        last_row = rows[-1]
        if (last_row.find('th') or 
            "sum" in last_row.get_text().lower() or 
            "total" in last_row.get_text().lower()):
            
            # Try to apply the table total style if it exists
            total_style_name = "= Table total"
            if total_style_name in [s.name for s in doc.styles]:
                for cell in table.rows[-1].cells:
                    for paragraph in cell.paragraphs:
                        try:
                            paragraph.style = total_style_name
                        except:
                            pass

def identify_list_blocks(lines):
    """
    Identify list blocks within markdown lines for later processing.
    
    Args:
        lines: List of markdown content lines
    
    Returns:
        list: List of tuple pairs (start_line, end_line) for each list block
    """
    list_blocks = []
    current_block = None
    in_code_block = False
    
    for i, line in enumerate(lines):
        # Check if this line starts or ends a code block
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            if current_block is not None:
                # This code block is part of the list item
                current_block.append(i)
            continue
            
        # If we're inside a code block, add to current list block if it exists
        if in_code_block and current_block is not None:
            current_block.append(i)
            continue
            
        # Check if line is a list item
        is_list_item = re.match(r'^(\s*)[-*+](.+)$', line) or re.match(r'^(\s*)\d+\.(.+)$', line)
        
        if is_list_item:
            if current_block is None:
                # Start a new list block
                current_block = [i]
            else:
                # Continue the existing block
                current_block.append(i)
        elif current_block is not None:
            if line.strip() == '':
                # Empty line - add to current block as it might be part of list formatting
                current_block.append(i)
            elif line.strip().startswith('```'):
                # Beginning of a code block - consider part of the list
                current_block.append(i)
            else:
                # Non-empty, non-list line after a list - end the block if it's not indented
                # Check if it's indented (part of a list item content)
                leading_space = len(line) - len(line.lstrip())
                if leading_space > 0:  # If indented, consider it part of the list
                    current_block.append(i)
                else:  # Not indented, end the list block
                    list_blocks.append((current_block[0], i-1))
                    current_block = None
    
    # Don't forget the last block if we're still in one
    if current_block is not None:
        list_blocks.append((current_block[0], len(lines)-1))
            
    return list_blocks

def process_inline_formatting(text):
    """
    Process inline formatting like bold, italic, etc.
    
    Args:
        text: Text to process
        
    Returns:
        BeautifulSoup element with formatted content
    """
    # Clean up any problematic markdown syntax first
    # Fix issues with inconsistent asterisks or markdown syntax
    text = re.sub(r'\*\*\*', '**', text)  # Replace *** with **
    text = re.sub(r'\\\*', '*', text)     # Replace \* with *
    text = re.sub(r'\*\*\*', '**', text)  # One more pass to catch any newly formed ***
    
    # Convert the text to HTML to handle the formatting
    html_snippet = markdown.markdown(text)
    
    # Parse the HTML
    soup = BeautifulSoup(html_snippet, 'html.parser')
    
    # If it's a simple paragraph with formatting, return the paragraph's HTML
    if soup.p:
        return soup.p
    
    # Otherwise, return the original text
    return text

def add_formatted_text_to_paragraph(paragraph, content):
    """
    Add formatted text to a paragraph.
    
    Args:
        paragraph: The paragraph to add text to
        content: The text content (string or BeautifulSoup element)
    """
    if isinstance(content, str):
        # Simple text, just add it
        paragraph.add_run(content)
    else:
        # It's a BeautifulSoup element, process its children
        for child in content.children:
            if isinstance(child, str):
                paragraph.add_run(child)
            elif child.name == 'strong' or child.name == 'b':
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name == 'em' or child.name == 'i':
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == 'code':
                run = paragraph.add_run(child.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                # For other elements, just get their text
                paragraph.add_run(child.get_text())

def process_text_with_markdown(text):
    """
    Process text with markdown formatting.
    
    Args:
        text: Text to process
        
    Returns:
        list: List of tuples (format_type, segment_text)
    """
    # Clean up problematic markdown syntax first
    # Fix issues with inconsistent asterisks and markdown syntax
    text = text.replace('\\*', '§ESCAPED_ASTERISK§')  # Temporarily protect escaped asterisks
    text = re.sub(r'\*{3,}', '**', text)  # Replace any *** or more with **
    text = re.sub(r'(\*{2})(\s*\*{1})|\*\s*\*\*', '**', text)  # Fix ** * or * ** patterns
    text = re.sub(r'(\*{1})(\s*\*{2})|\*\*\s*\*', '**', text)  # Fix * ** or ** * patterns
    text = text.replace('§ESCAPED_ASTERISK§', '*')  # Restore escaped asterisks
    
    # Process the original text to find all markdown formatted sections
    segments = []
    
    # First, handle code backticks to prevent issues with other patterns
    code_pattern = r'`(.*?)`'
    code_segments = []
    last_end = 0
    
    for match in re.finditer(code_pattern, text):
        if match.start() > last_end:
            code_segments.append(('normal', text[last_end:match.start()]))
        code_segments.append(('code', match.group(1)))
        last_end = match.end()
    
    if last_end < len(text):
        code_segments.append(('normal', text[last_end:]))
    
    # Now process bold and italic in each non-code segment
    for segment_type, segment_text in code_segments:
        if segment_type == 'code':
            segments.append(('code', segment_text))
        else:
            # Process bold (must be done before italic to handle nested formatting)
            bold_pattern = r'\*\*(.*?)\*\*|__(.*?)__'
            bold_segments = []
            last_end = 0
            
            for match in re.finditer(bold_pattern, segment_text):
                if match.start() > last_end:
                    bold_segments.append(('normal', segment_text[last_end:match.start()]))
                
                # Add the bold text (without ** or __)
                bold_text = match.group(1) if match.group(1) is not None else match.group(2)
                bold_segments.append(('bold', bold_text))
                
                last_end = match.end()
            
            if last_end < len(segment_text):
                bold_segments.append(('normal', segment_text[last_end:]))
            
            # Process italic in each non-bold segment
            for bold_type, bold_content in bold_segments:
                if bold_type == 'bold':
                    # For bold text, check for nested italic
                    italic_pattern = r'\*(.*?)\*|_(.*?)_'
                    italic_segments = []
                    last_end = 0
                    
                    for match in re.finditer(italic_pattern, bold_content):
                        if match.start() > last_end:
                            italic_segments.append(('bold', bold_content[last_end:match.start()]))
                        
                        # Add the italic text (without * or _)
                        italic_text = match.group(1) if match.group(1) is not None else match.group(2)
                        italic_segments.append(('bold+italic', italic_text))
                        
                        last_end = match.end()
                    
                    if last_end < len(bold_content):
                        italic_segments.append(('bold', bold_content[last_end:]))
                    
                    segments.extend(italic_segments)
                else:
                    # For normal text, check for italic
                    italic_pattern = r'\*(.*?)\*|_(.*?)_'
                    italic_segments = []
                    last_end = 0
                    
                    for match in re.finditer(italic_pattern, bold_content):
                        if match.start() > last_end:
                            italic_segments.append(('normal', bold_content[last_end:match.start()]))
                        
                        # Add the italic text (without * or _)
                        italic_text = match.group(1) if match.group(1) is not None else match.group(2)
                        italic_segments.append(('italic', italic_text))
                        
                        last_end = match.end()
                    
                    if last_end < len(bold_content):
                        italic_segments.append(('normal', bold_content[last_end:]))
                    
                    segments.extend(italic_segments)
    
    return segments

def clear_template_body_except_top(doc):
    """
    Clear out most of the body content from the template document, 
    but preserve the header/logo section at the top.
    This approach is similar to the original working code where we
    only remove paragraphs but leave tables intact.
    
    Args:
        doc: The python-docx Document object
    """
    # Access the body element
    body = doc._body
    
    # Identify any tables in the document
    tables_found = 0
    table_indices = []
    
    # First, scan the document to identify tables
    for i, element in enumerate(body._element):
        if element.tag.endswith('tbl'):
            tables_found += 1
            table_indices.append(i)
            
    print(f"Found {tables_found} tables in template")
    
    # Create a list of elements to remove:
    # 1. Preserve the first table (logo)
    # 2. Remove all paragraphs except those in the header/footer
    # 3. Remove all other tables after the first one
    
    elements_to_remove = []
    for i, element in enumerate(body._element[:]):
        # Skip the first table - it's the logo section
        if i in table_indices and i == table_indices[0]:
            continue
            
        # Remove all other paragraphs and tables
        if element.tag.endswith('p') or (element.tag.endswith('tbl') and i != table_indices[0]):
            elements_to_remove.append(element)
    
    # Remove the identified elements
    for element in elements_to_remove:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
    
    print(f"Template body cleared - preserved top section, removed {len(elements_to_remove)} elements")

def clear_template_body_except_header(doc):
    """
    Clear out the body content from the template document, 
    but preserve the CBS logo and header information at the top.
    
    Args:
        doc: The python-docx Document object
    """
    # First identify the elements we need to keep - typically these are at the start of the document
    keep_element_count = 0
    logo_found = False
    details_found = False
    
    # Access the body element
    body = doc._body
    all_elements = list(body._element[:])
    
    # We'll search through the first few elements to find what we need to keep
    for i, element in enumerate(all_elements[:10]):  # Look through first 10 elements
        element_text = element.text if hasattr(element, 'text') else ""
        
        # Check if this is a table that might contain the CBS logo
        if element.tag.endswith('tbl'):
            # This is likely the logo table - keep it
            logo_found = True
            keep_element_count = i + 1
        
        # Look for the "DETAILS" section or a similar header
        if "DETAILS" in element_text:
            details_found = True
            keep_element_count = i + 1
        
        # If we found both the logo and details, we can stop searching
        if logo_found and details_found:
            break
    
    # If we didn't find these elements, try a different approach
    # Just preserve the first table and a couple of paragraphs
    if not logo_found and not details_found and len(all_elements) > 3:
        for i, element in enumerate(all_elements[:5]):
            if element.tag.endswith('tbl'):
                keep_element_count = i + 2  # Keep the table and the next element
                break
    
    # Ensure we're keeping at least one element
    keep_element_count = max(keep_element_count, 1)
    
    # Now remove all elements except those we want to keep
    elements_to_remove = []
    for i, element in enumerate(all_elements):
        # Skip elements we want to keep
        if i < keep_element_count:
            continue
        
        # Remove paragraphs and tables after the keep limit
        if element.tag.endswith('tbl') or element.tag.endswith('p'):
            elements_to_remove.append(element)
    
    # Remove the unwanted elements
    for element in elements_to_remove:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
    
    print(f"Template body cleared - kept {keep_element_count} elements at the top, removed {len(elements_to_remove)} elements")

def check_document_sections(doc):
    """
    Debug function to check if headers and footers are present in the document.
    
    Args:
        doc: The python-docx Document object
    """
    try:
        sections = doc.sections
        print(f"Document has {len(sections)} section(s)")
        
        for i, section in enumerate(sections):
            print(f"Section {i+1}:")
            
            # Check for headers
            if section.header:
                header_paragraphs = len(section.header.paragraphs)
                print(f"  - Header: {header_paragraphs} paragraph(s)")
            else:
                print("  - No header")
                
            # Check for footers
            if section.footer:
                footer_paragraphs = len(section.footer.paragraphs)
                print(f"  - Footer: {footer_paragraphs} paragraph(s)")
            else:
                print("  - No footer")
    except Exception as e:
        print(f"Error checking sections: {str(e)}")

def create_nested_list_items_with_styles(doc, items, styles_info):
    """
    Create list items with correct nesting based on indentation and using template styles.
    
    Args:
        doc: The docx Document object
        items: List of tuples (indent, content, is_ordered, line_num)
        styles_info: Dictionary containing style information from the template
    """
    # Use bullet characters from template, or fallback to defaults
    # bullet_chars = styles_info.get('default_bullets', ['•', '◦', '▪', '▫'])
    
    # Get style names for bullet and numbered lists
    # For numbered lists, specifically look for styles containing "Text numbering"
    bullet_styles = list(styles_info.get('bullet_styles', {}).keys()) or ['–bullet 1', '–bullet 2', '–bullet 3', '–bullet 4']
    
    # For Fix 2: Look specifically for text numbering styles and prioritize them
    text_numbering_styles = []
    for style_name in styles_info.get('paragraph_styles', {}):
        if 'text numbering' in style_name.lower() or 'textnumbering' in style_name.lower():
            text_numbering_styles.append(style_name)
    
    # If no specific text numbering styles found, use fallbacks
    numbering_styles = text_numbering_styles or ['| Text numbering 1', '| Text numbering a']
    
    # Print available numbering styles for debugging
    print(f"Available numbering styles: {numbering_styles}")
    
    # Track item counters for ordered lists at each nesting level
    number_counters = {}
    
    # Compute minimum indentation to establish base level
    min_indent = min(item[0] for item in items) if items else 0
    
    # Process each item
    for indent, content, is_ordered, line_num in items:
        # Calculate actual nesting level (relative to minimum indent)
        level = (indent - min_indent) // 2
        level = min(level, 3)  # Cap level to reasonable depth
        
        # Create paragraph 
        p = doc.add_paragraph()
        
        if is_ordered:
            # Select appropriate numbering style for this level
            style_index = min(level, len(numbering_styles) - 1)
            try:
                style_name = numbering_styles[style_index]
                p.style = style_name
                print(f"Applied numbering style: {style_name}")
            except Exception as e:
                print(f"Failed to apply numbering style: {e}")
                p.style = 'Normal'
            
            # For Fix 2: Don't add the number prefix manually for properly styled lists
            # Let the paragraph style handle the numbering
        else:
            # Select bullet style for this level
            style_index = min(level, len(bullet_styles) - 1)
            try:
                p.style = bullet_styles[style_index]
            except:
                p.style = 'Normal'
        
        # Add content with formatting support
        add_formatted_text_to_paragraph(p, content)

def process_list_block_with_styles(doc, lines, start_line, end_line, styles_info):
    """
    Parse and process a list block using template styles, handling nested and multi-level lists.
    
    Args:
        doc: The docx Document object
        lines: List of markdown lines
        start_line: Start index of the list block
        end_line: End index of the list block
        styles_info: Dictionary containing style information from the template
        
    Returns:
        bool: True if any list items were processed
    """
    # Extract all list items with their indentation levels
    items = []
    current_list_item = None
    in_code_block = False
    code_block_lines = []
    
    i = start_line
    while i <= end_line:
        line = lines[i].rstrip()
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_start = i
                language = line.strip()[3:]
            else:
                in_code_block = False
                code_text = '\n'.join(code_block_lines)
                add_code_block(doc, code_text, language)
                code_block_lines = []
            
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Check for list items (both unordered and ordered)
        ul_match = re.match(r'^(\s*)[-*+](.+)$', line)
        ol_match = re.match(r'^(\s*)\d+\.(.+)$', line)
        
        if ul_match or ol_match:
            # If we had a previous list item, add it now
            if current_list_item is not None:
                items.append(current_list_item)
            
            match = ul_match or ol_match
            indent = len(match.group(1))
            content = match.group(2).strip()
            is_ordered = bool(ol_match)
            
            # Process inline formatting in list items
            content = process_inline_formatting(content)
            
            # Store this list item
            current_list_item = (indent, content, is_ordered, i)
        
        i += 1
    
    # Add the last list item if any
    if current_list_item is not None:
        items.append(current_list_item)
    
    # Process items with correct nesting
    if items:
        create_nested_list_items_with_styles(doc, items, styles_info)
        return True
    
    return False

def create_docx_from_readme_with_styles(readme_path, template_path, output_path):
    """
    Convert a Markdown file to a Word document using a template with style extraction.
    
    Args:
        readme_path (str): Path to the Markdown file
        template_path (str): Path to the Word template
        output_path (str): Path to save the output Word document
    
    Returns:
        bool: True if conversion was successful, False otherwise
    """
    try:
        # Load the template document
        doc = Document(template_path)
        
        # Extract style information from the template
        styles_info = extract_template_styles(doc)
        
        # Print available styles for debugging
        print("Available styles in template:")
        for style_name in styles_info['paragraph_styles']:
            print(f"- {style_name}")
        
        # Use our new function to preserve the top section while clearing the rest
        clear_template_body_except_top(doc)

        
        # Debug: Check headers and footers after clearing
        print("After clearing template body:")
        check_document_sections(doc)
        
        # Preserve headers and footers by removing only body paragraphs
        body = doc._body
        for element in body._element[:]:
            if element.tag.endswith('p'):
                body._element.remove(element)

        # Read the README file and remove soft carriage returns
        with open(readme_path, 'r', encoding='utf-8') as f:
            readme_content = f.read().replace('\r\n', '\n').replace('\r', '\n')
        
        # Split the markdown content into lines for direct parsing
        markdown_lines = readme_content.split('\n')
        
        # Convert Markdown to HTML for non-list elements
        html_content = markdown.markdown(readme_content, extensions=['extra', 'tables', 'fenced_code'])
        soup = BeautifulSoup(html_content, 'html.parser')

        # Find all list blocks in the markdown
        list_blocks = identify_list_blocks(markdown_lines)
        
        # Create a map to find list blocks by line
        list_block_map = {}
        for start, end in list_blocks:
            for i in range(start, end + 1):
                list_block_map[i] = (start, end)
        
        # Keep track of whether we've seen the H1 heading
        h1_heading_processed = False
        
        # Track lines that have been processed
        processed_lines = set()
        
        # Track processed list blocks to avoid duplicates
        processed_list_blocks = set()
        
        # Create an additional set to track which paragraphs have inline lists processed
        inline_list_paragraphs = set()
        
        # ===== PROCESS THE DOCUMENT LINE BY LINE TO MAINTAIN ORDER =====
        i = 0
        while i < len(markdown_lines):
            line = markdown_lines[i].strip()
            
            # Skip empty lines and horizontal rules (---, ___, ***)
            if not line or re.match(r'^[-_*]{3,}$', line):
                processed_lines.add(i)
                i += 1
                continue
            
            # Check if it's a list block
            if i in list_block_map:
                start, end = list_block_map[i]
                
                # Skip if we already processed this list block
                if (start, end) not in processed_list_blocks:
                    process_list_block_with_styles(doc, markdown_lines, start, end, styles_info)
                    processed_list_blocks.add((start, end))
                    
                    # Mark all lines in this block as processed
                    for j in range(start, end + 1):
                        processed_lines.add(j)
                
                # Skip to the end of the block
                i = end + 1
                continue
            
            # Check if it's a heading
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                
                # Adjust heading levels
                if level == 1:
                    # First # becomes Heading 0
                    style_name = 'Heading 0'
                else:
                    # Subsequent levels are shifted down by 1
                    style_name = f'Heading {level - 1}'
                
                # Create heading paragraph
                p = doc.add_paragraph()
                try:
                    p.style = style_name
                except KeyError:
                    print(f"Warning: Style '{style_name}' not found. Using 'Normal'.")
                    p.style = 'Normal'
                
                # Clean any manual numbering
                text = clean_heading_text(text)
                
                # Process any markdown formatting in headings
                segments = process_text_with_markdown(text)
                for format_type, segment_text in segments:
                    run = p.add_run(segment_text)
                    if format_type == 'bold' or format_type == 'bold+italic':
                        run.bold = True
                    if format_type == 'italic' or format_type == 'bold+italic':
                        run.italic = True
                    if format_type == 'code':
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                
                processed_lines.add(i)
                i += 1
                continue
            
            # Check if it's a code block
            if line.startswith('```'):
                # Extract language if specified
                language = line[3:].strip()
                
                # Collect code lines until closing ```
                code_lines = []
                i += 1
                start_code_block = i - 1
                
                while i < len(markdown_lines) and not markdown_lines[i].strip().startswith('```'):
                    code_lines.append(markdown_lines[i])
                    processed_lines.add(i)
                    i += 1
                
                # Skip the closing ```
                end_code_block = i
                processed_lines.add(i)
                i += 1
                
                # Add code block
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    add_code_block(doc, code_text, language)
                
                continue
            
            # If it's a table, use our enhanced table handler
            if line.startswith('|') and i + 1 < len(markdown_lines) and markdown_lines[i + 1].strip().startswith('|'):
                # Find table in HTML
                tables = soup.find_all('table')
                table_start = i
                
                # Skip ahead to find the end of the table
                while i < len(markdown_lines) and (markdown_lines[i].strip().startswith('|') or markdown_lines[i].strip() == ''):
                    processed_lines.add(i)
                    i += 1
                table_end = i - 1
                
                # Process the table only once
                for table in tables:
                    # Add the table with our enhanced function
                    table_text = ''.join(markdown_lines[table_start:table_end+1])
                    add_table_with_styles(doc, table, styles_info)
                    break
                    
                continue
            
            # Only process each line once to avoid duplication
            if i in processed_lines:
                i += 1
                continue
                
            # It's a regular paragraph, find in HTML to preserve formatting
            paragraphs = soup.find_all('p')
            plain_text_map = {}
            for p in paragraphs:
                plain_text = p.get_text().strip()
                plain_text_map[plain_text] = p
            
            # Try to find closest match
            matched = False
            for plain_text, p_element in plain_text_map.items():
                if line in plain_text or plain_text in line:
                    # Check if this paragraph contains a list item that will be processed separately
                    is_list_item = False
                    for list_start, list_end in list_blocks:
                        for list_line in range(list_start, list_end + 1):
                            list_text = markdown_lines[list_line].strip()
                            if list_text and (list_text in plain_text or plain_text in list_text):
                                is_list_item = True
                                break
                        if is_list_item:
                            break
                
                    # Skip if it's a list item
                    if not is_list_item:
                        add_paragraph(doc, p_element, 'Normal')
                        matched = True
                    break
            
            if not matched:
                # Skip the line if it appears to be a formatting artifact 
                # like a single dash or lone asterisk, etc.
                if line == '--' or line == '*' or line == '---':
                    processed_lines.add(i)
                    
                    # For horizontal rules, add a proper horizontal line
                    if line == '---':
                        hr_paragraph = doc.add_paragraph()
                        hr_paragraph.style = 'Normal'
                        # Add a thin horizontal line by adding a bottom border to the paragraph
                        hr_paragraph.paragraph_format.bottom_border.width = Pt(1)
                        
                    i += 1
                    continue
                    
                # Special handling for contact information at the end of the document
                if line.startswith('*For') and line.endswith('*'):
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    run = p.add_run(line.strip('*'))
                    run.italic = True
                    processed_lines.add(i)
                    i += 1
                    continue
                    
                # Skip line if it appears to be part of a list that will be processed separately
                is_list_line = False
                for list_start, list_end in list_blocks:
                    if list_start <= i <= list_end:
                        is_list_line = True
                        break
                
                if not is_list_line:
                    # If no match found, try to process markdown formatting directly
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    
                    # Fix special cases for contact info/footnotes
                    if line.startswith('*For') and line.endswith('*'):
                        # Handle the contact info line as italic
                        run = p.add_run(line.strip('*'))
                        run.italic = True
                    else:
                        # Process markdown formatting
                        segments = process_text_with_markdown(line)
                        for format_type, segment_text in segments:
                            run = p.add_run(segment_text)
                            if format_type == 'bold' or format_type == 'bold+italic':
                                run.bold = True
                            if format_type == 'italic' or format_type == 'bold+italic':
                                run.italic = True
                            if format_type == 'code':
                                run.font.name = 'Courier New'
                                run.font.size = Pt(10)
            
            processed_lines.add(i)
            i += 1
        
        # Save the document
        doc.save(output_path)
        print(f"Document saved to {output_path}")
        return True
        
    except Exception as e:
        import traceback
        print(f"Error during conversion: {str(e)}")
        print(traceback.format_exc())
        return False

def convert_file_with_styles(md_file_path, template_path, output_path=None):
    """
    Convert a single Markdown file to DOCX with style extraction
    
    Args:
        md_file_path (str): Path to the markdown file
        template_path (str): Path to the Word template
        output_path (str, optional): Path to save the output file. If None, will use the same name as the input file.
    
    Returns:
        tuple: (success, output_path)
    """
    try:
        # If output_path is not specified, create one based on input file
        if output_path is None:
            output_dir = os.path.dirname(md_file_path)
            file_name = os.path.splitext(os.path.basename(md_file_path))[0]
            output_path = os.path.join(output_dir, f"{file_name}.docx")
        
        # If output_path is a directory, append file name
        elif os.path.isdir(output_path):
            file_name = os.path.splitext(os.path.basename(md_file_path))[0]
            output_path = os.path.join(output_path, f"{file_name}.docx")
            
        # Create parent directory if it doesn't exist
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        
        # Call conversion function
        success = create_docx_from_readme_with_styles(md_file_path, template_path, output_path)
        
        return success, output_path
    except Exception as e:
        import traceback
        print(f"Error converting file: {str(e)}")
        print(traceback.format_exc())
        return False, None

def convert_folder_with_styles(folder_path, template_path, output_folder=None):
    """
    Convert all Markdown files in a folder to DOCX with style extraction
    
    Args:
        folder_path (str): Path to the folder containing markdown files
        template_path (str): Path to the Word template
        output_folder (str, optional): Path to save the output files. If None, will use the same folder as the input.
    
    Returns:
        list: List of (success, output_path) tuples for each converted file
    """
    try:
        # If output_folder is not specified, use the input folder
        if output_folder is None:
            output_folder = folder_path
        
        # Make sure the output folder exists
        os.makedirs(output_folder, exist_ok=True)
        
        results = []
        
        # Find all markdown files in the folder
        md_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.md')]
        
        # Convert each file
        for md_file in md_files:
            md_file_path = os.path.join(folder_path, md_file)
            file_name = os.path.splitext(md_file)[0]
            output_path = os.path.join(output_folder, f"{file_name}.docx")
            
            # Call the convert_file function to handle the conversion
            success, actual_output_path = convert_file_with_styles(md_file_path, template_path, output_path)
            results.append((success, actual_output_path if success else output_path))
        
        return results
    except Exception as e:
        import traceback
        print(f"Error converting folder: {str(e)}")
        print(traceback.format_exc())
        return [(False, None)]

# Original functions (keep for backward compatibility)

def create_docx_from_readme(readme_path, template_path, output_path):
    """
    Convert a Markdown file to a Word document using a template.
    (Original function - maintained for backward compatibility)
    
    Args:
        readme_path (str): Path to the Markdown file
        template_path (str): Path to the Word template
        output_path (str): Path to save the output Word document
    
    Returns:
        bool: True if conversion was successful, False otherwise
    """
    # Call the enhanced version with style extraction
    return create_docx_from_readme_with_styles(readme_path, template_path, output_path)

def convert_file(md_file_path, template_path, output_path=None):
    """
    Convert a single Markdown file to DOCX
    (Original function - maintained for backward compatibility)
    
    Args:
        md_file_path (str): Path to the markdown file
        template_path (str): Path to the Word template
        output_path (str, optional): Path to save the output file. If None, will use the same name as the input file.
    
    Returns:
        tuple: (success, output_path)
    """
    # Call the enhanced version with style extraction
    return convert_file_with_styles(md_file_path, template_path, output_path)

def convert_folder(folder_path, template_path, output_folder=None):
    """
    Convert all Markdown files in a folder to DOCX
    (Original function - maintained for backward compatibility)
    
    Args:
        folder_path (str): Path to the folder containing markdown files
        template_path (str): Path to the Word template
        output_folder (str, optional): Path to save the output files. If None, will use the same folder as the input.
    
    Returns:
        list: List of (success, output_path) tuples for each converted file
    """
    # Call the enhanced version with style extraction
    return convert_folder_with_styles(folder_path, template_path, output_folder)